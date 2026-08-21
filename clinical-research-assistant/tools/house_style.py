#!/usr/bin/env python3
"""
house_style.py — mechanical enforcement of the CRA house document format (lesson L070).

The standard, from skills/references/writing-style.md
"House Document Formatting — Universal Standard":

  * Tables are BLACK AND WHITE. No cell fills, no coloured text, no coloured borders.
    Structure is carried by rules and whitespace; emphasis by bold weight only.
  * The default font is TIMES NEW ROMAN everywhere, including figures.
    Fallback chain: Times New Roman -> Liberation Serif -> generic serif.
  * L051 bolding (p<0.05 in Table_1; BH-FDR q<0.05 elsewhere) is preserved — bold is
    monochrome emphasis and remains the correct significance flag.
  * Figures keep colour ONLY where colour encodes a variable, and must stay
    greyscale-separable.

Usage
-----
    python3 tools/house_style.py FILE [FILE ...]        # fix in place
    python3 tools/house_style.py --check FILE [...]     # report only, exit 1 on violation
    python3 tools/house_style.py --reference-doc OUT    # emit a pandoc reference .docx

In matplotlib, call apply_matplotlib() before plotting, or pass the rcParams dict.

Companion to tools/voice_check.py (prose voice); this one governs format.
"""
from __future__ import annotations
import sys, os

FONT = "Times New Roman"
FALLBACKS = [FONT, "Liberation Serif", "Nimbus Roman", "DejaVu Serif", "serif"]
_WHITE = {None, "00000000", "FFFFFFFF", "FFFFFF", "00FFFFFF"}


# ------------------------------------------------------------------ Word
def enforce_docx(path: str, check: bool = False) -> list[str]:
    from docx import Document
    from docx.shared import RGBColor
    from docx.oxml.ns import qn
    v: list[str] = []
    doc = Document(path)

    def fix_run(run, where):
        if (run.font.name or "") != FONT:
            v.append(f"{where}: font {run.font.name!r}")
            if not check:
                run.font.name = FONT
                rPr = run._element.get_or_add_rPr()
                rF = rPr.find(qn("w:rFonts"))
                if rF is None:
                    rF = rPr.makeelement(qn("w:rFonts"), {}); rPr.append(rF)
                for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                    rF.set(qn(slot), FONT)
        col = run.font.color
        if col is not None and col.rgb is not None and str(col.rgb) != "000000":
            v.append(f"{where}: coloured text #{col.rgb}")
            if not check:
                run.font.color.rgb = RGBColor(0, 0, 0)

    for st_name in ("Normal", "Table Grid", "Heading 1", "Heading 2", "Heading 3", "Title"):
        try:
            st = doc.styles[st_name]
        except KeyError:
            continue
        if (st.font.name or "") != FONT:
            v.append(f"style {st_name}: font {st.font.name!r}")
            if not check:
                st.font.name = FONT
                rPr = st.element.get_or_add_rPr()
                rF = rPr.find(qn("w:rFonts"))
                if rF is None:
                    rF = rPr.makeelement(qn("w:rFonts"), {}); rPr.append(rF)
                for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                    rF.set(qn(slot), FONT)

    for i, p in enumerate(doc.paragraphs):
        for r in p.runs:
            fix_run(r, f"para {i}")

    for ti, t in enumerate(doc.tables):
        if t.style is not None and t.style.name not in ("Table Grid", "Normal Table"):
            v.append(f"table {ti}: style {t.style.name!r} (not monochrome)")
            if not check:
                t.style = "Table Grid"
        for row in t.rows:
            for cell in row.cells:
                # strip any cell shading
                tcPr = cell._tc.get_or_add_tcPr()
                for sh in tcPr.findall(qn("w:shd")):
                    fill = sh.get(qn("w:fill"))
                    if fill and fill.upper() not in ("AUTO", "FFFFFF"):
                        v.append(f"table {ti}: cell shading #{fill}")
                    if not check:
                        tcPr.remove(sh)
                for p in cell.paragraphs:
                    # drop empty runs — invisible, but they inherit the template font
                    for r in list(p.runs):
                        if r.text == "":
                            if not check:
                                r._element.getparent().remove(r._element)
                            continue
                        fix_run(r, f"table {ti}")
    if not check:
        doc.save(path)
    return v


# ----------------------------------------------------------------- Excel
def enforce_xlsx(path: str, check: bool = False) -> list[str]:
    from openpyxl import load_workbook
    from openpyxl.styles import Font, PatternFill
    v: list[str] = []
    wb = load_workbook(path)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for c in row:
                if c.value is None:
                    continue
                f = c.fill
                if f is not None and f.fill_type == "solid" and \
                        getattr(f.fgColor, "rgb", None) not in _WHITE:
                    v.append(f"{ws.title}!{c.coordinate}: fill #{f.fgColor.rgb}")
                    if not check:
                        c.fill = PatternFill(fill_type=None)
                fn = c.font
                if (fn.name or "") != FONT or (fn.color is not None and
                        getattr(fn.color, "rgb", None) not in _WHITE | {"FF000000", "000000"}):
                    if (fn.name or "") != FONT:
                        v.append(f"{ws.title}!{c.coordinate}: font {fn.name!r}")
                    if not check:
                        # preserve bold/italic/size — L051 bolding must survive
                        c.font = Font(name=FONT, bold=fn.bold, italic=fn.italic,
                                      size=fn.size, underline=fn.underline)
    if not check:
        wb.save(path)
    return v


# ------------------------------------------------------------ matplotlib
def matplotlib_rc() -> dict:
    return {"font.family": "serif", "font.serif": FALLBACKS,
            "mathtext.fontset": "dejavuserif"}


def apply_matplotlib():
    """Set Times New Roman for all subsequent matplotlib output. Returns the resolved family."""
    import matplotlib
    from matplotlib import font_manager
    matplotlib.rcParams.update(matplotlib_rc())
    avail = {f.name for f in font_manager.fontManager.ttflist}
    for cand in FALLBACKS:
        if cand in avail:
            return cand
    return "serif (no named fallback resolved)"


# -------------------------------------------------------- reference .docx
def make_reference_doc(out: str):
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    d = Document()
    for name in ("Normal", "Table Grid", "Heading 1", "Heading 2", "Heading 3", "Title"):
        try:
            st = d.styles[name]
        except KeyError:
            continue
        st.font.name = FONT
        if name == "Normal":
            st.font.size = Pt(11)
        rPr = st.element.get_or_add_rPr()
        rF = rPr.find(qn("w:rFonts"))
        if rF is None:
            rF = rPr.makeelement(qn("w:rFonts"), {}); rPr.append(rF)
        for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rF.set(qn(slot), FONT)
    d.save(out)
    return out


def main(argv):
    args = [a for a in argv[1:] if not a.startswith("--")]
    check = "--check" in argv
    if "--reference-doc" in argv:
        print("wrote", make_reference_doc(args[0])); return 0
    if not args:
        print(__doc__); return 0
    total = 0
    for p in args:
        ext = os.path.splitext(p)[1].lower()
        if ext == ".docx":
            v = enforce_docx(p, check)
        elif ext == ".xlsx":
            v = enforce_xlsx(p, check)
        else:
            print(f"skip (unsupported): {p}"); continue
        total += len(v)
        verb = "violations" if check else "fixed"
        print(f"{os.path.basename(p)}: {len(v)} {verb}")
        for x in v[:6]:
            print(f"    {x}")
        if len(v) > 6:
            print(f"    ... and {len(v)-6} more")
    return 1 if (check and total) else 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
